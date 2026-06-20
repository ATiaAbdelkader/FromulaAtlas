import docx
from docx import Document

doc = Document('/home/z/my-project/upload/Agri_Formulas_Metrics_Handbook_v3_Complete.docx')

print(f"=== TOTAL PARAGRAPHS: {len(doc.paragraphs)} ===")
print(f"=== TOTAL TABLES: {len(doc.tables)} ===")
print()

# Print all paragraphs with style
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else "None"
        print(f"[P{i}][{style}] {para.text}")

print("\n\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
