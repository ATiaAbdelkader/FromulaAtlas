"""
Parse the Agri-Formulas Handbook docx into a structured JSON for the web app.

Each formula table is 3 rows:
- Row 0: Title "X.X  Name" (both cells identical)
- Row 1: Formula (both cells identical)
- Row 2: Left cell = "Variables\n<vars>"; Right cell = "Purpose\n<text>\nExample\n<text>\n⚠ Pitfall: <text>"
"""

import re
import json
from docx import Document
from docx.oxml.ns import qn

DOC_PATH = '/home/z/my-project/upload/Agri_Formulas_Metrics_Handbook_v3_Complete.docx'
OUT_PATH = '/home/z/my-project/scripts/agri_formulas.json'


def clean(text):
    if text is None:
        return ''
    return text.strip()


def detect_part_number(text):
    m = re.match(r'^\s*Part\s+([IVXLCDM]+)\s*$', text, re.IGNORECASE)
    return m.group(1) if m else None


def parse_chapter_title(text):
    m = re.match(r'^\s*Chapter\s+(\d+)\s+(.+?)\s*$', text)
    if m:
        return int(m.group(1)), m.group(2).strip()
    return None


def split_labeled_block(text):
    """
    Split a block of text into labeled sections.
    Labels we care about: Variables, Purpose, Example, Pitfall, Decision, Note, Where, Units
    Each label starts a new section; lines following until the next label are the section content.
    Lines beginning with ⚠ or ▲ or ♻ are also treated as label-starts.
    Returns dict: {label_lower: content_string}
    """
    if not text:
        return {}
    lines = text.split('\n')
    result = {}
    current_label = None
    current_buf = []

    LABELS = ['variables', 'variable', 'purpose', 'example', 'pitfall',
              'decision', 'note', 'where', 'units', 'unit', 'interpretation',
              'threshold', 'use', 'output', 'benchmark']

    def flush():
        nonlocal current_label, current_buf
        if current_label:
            content = '\n'.join(current_buf).strip()
            # Map label to canonical key
            ll = current_label.lower()
            # Strip trailing ':' or whitespace
            ll = ll.rstrip(':').strip()
            # Find the canonical key
            canonical = None
            for L in LABELS:
                if ll == L or ll.startswith(L):
                    canonical = L
                    break
            if canonical is None:
                # If unknown, but starts with a symbol
                if current_label.startswith('⚠'):
                    canonical = 'pitfall'
                elif current_label.startswith('▲'):
                    canonical = 'decision'
                elif current_label.startswith('♻'):
                    canonical = 'note'
                else:
                    canonical = 'extra'
            if canonical in result:
                result[canonical] += '\n' + content
            else:
                result[canonical] = content
        current_label = None
        current_buf = []

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            # Empty line - keep as paragraph break in current buffer
            if current_buf:
                current_buf.append('')
            continue

        # Check if line is a label (e.g., "Variables", "Purpose", "Example", "⚠ Pitfall: ...")
        is_label = False
        label_text = ''
        label_content_after = ''

        # Symbol-prefixed labels
        if line_stripped.startswith('⚠'):
            is_label = True
            label_text = 'pitfall'
            label_content_after = line_stripped.lstrip('⚠').lstrip(': ').strip()
        elif line_stripped.startswith('▲'):
            is_label = True
            label_text = 'decision'
            label_content_after = line_stripped.lstrip('▲').lstrip(': ').strip()
        elif line_stripped.startswith('♻'):
            is_label = True
            label_text = 'note'
            label_content_after = line_stripped.lstrip('♻').lstrip(': ').strip()
        else:
            # Word label: try to match "Label:" or "Label" at start
            m = re.match(r'^([A-Za-z][A-Za-z\s]{0,30}?)(?:\s*[:：]\s*|\s+)(.*)$', line_stripped)
            if m:
                potential_label = m.group(1).strip().lower()
                # Check if it's a known label
                for L in LABELS:
                    if potential_label == L:
                        is_label = True
                        label_text = L
                        label_content_after = m.group(2).strip()
                        break
                # Also handle "Pitfall:" directly
                if not is_label and potential_label.startswith('pitfall'):
                    is_label = True
                    label_text = 'pitfall'
                    label_content_after = m.group(2).strip()

        if is_label:
            flush()
            current_label = label_text
            if label_content_after:
                current_buf.append(label_content_after)
        else:
            if current_label is None:
                # No current label - try to assign based on context or skip
                # This can happen for the very first line if no label
                current_label = 'extra'
                current_buf = []
            current_buf.append(line_stripped)

    flush()
    return result


def parse_formula_table(table):
    """Parse a 3-row formula table into a formula dict."""
    rows = table.rows
    if len(rows) < 2:
        return None

    # Row 0: title
    title_cell = clean(rows[0].cells[0].text)
    m = re.match(r'^(\d+(?:\.\d+[a-z]?)?)\s+(.+?)\s*$', title_cell)
    if m:
        code = m.group(1)
        name = m.group(2).strip()
    else:
        code = ''
        name = title_cell

    # Row 1: formula
    formula_text = clean(rows[1].cells[0].text)

    # Row 2: left = variables block; right = purpose/example/pitfall block
    variables = ''
    right_block = ''
    if len(rows) >= 3:
        left = clean(rows[2].cells[0].text)
        right = clean(rows[2].cells[1].text) if len(rows[2].cells) > 1 else ''
        variables = left
        right_block = right

    # Parse the right block into labeled sections
    right_parsed = split_labeled_block(right_block)

    # Parse the variables block (it starts with "Variables" label)
    vars_parsed = split_labeled_block(variables)

    # Combine
    final_vars = vars_parsed.get('variables', '') or vars_parsed.get('extra', '')
    # If variables block didn't have "Variables" label, just use raw text minus the label
    if not final_vars and variables:
        # Strip leading "Variables" word if present
        v = re.sub(r'^[Vv]ariables\s*[:：]?\s*', '', variables)
        final_vars = v

    purpose = right_parsed.get('purpose', '')
    example = right_parsed.get('example', '')
    pitfall = right_parsed.get('pitfall', '')
    decision = right_parsed.get('decision', '')
    notes = right_parsed.get('note', '')

    # Sometimes pitfall is part of example block - check
    if not pitfall and example:
        # Look for "⚠ Pitfall:" inside example
        pm = re.search(r'⚠\s*Pitfall\s*[:：]\s*(.+?)$', example, re.IGNORECASE | re.DOTALL)
        if pm:
            pitfall = pm.group(1).strip()
            example = re.sub(r'⚠\s*Pitfall\s*[:：]\s*.+?$', '', example, flags=re.IGNORECASE | re.DOTALL).strip()

    # Sometimes variables also contain a "Where:" label
    # Sometimes purpose and example are concatenated without labels
    if not purpose and right_block:
        # Try to split right block at "Example" boundary
        parts = re.split(r'\n\s*Example\s*[:：]?\s*', right_block, maxsplit=1, flags=re.IGNORECASE)
        if len(parts) == 2:
            purpose = parts[0].replace('Purpose', '', 1).strip()
            # The remainder may contain example + pitfall
            remainder = parts[1]
            pm = re.search(r'⚠\s*Pitfall\s*[:：]\s*(.+?)$', remainder, re.IGNORECASE | re.DOTALL)
            if pm:
                pitfall = pm.group(1).strip()
                example = re.sub(r'⚠\s*Pitfall\s*[:：]\s*.+?$', '', remainder, flags=re.IGNORECASE | re.DOTALL).strip()
            else:
                example = remainder.strip()

    # Clean up: strip leading "Variables" label from final_vars
    final_vars = re.sub(r'^[Vv]ariables\s*[:：]?\s*\n?', '', final_vars).strip()

    # Strip leading "Purpose" / "Example" labels if they leaked through
    purpose = re.sub(r'^[Pp]urpose\s*[:：]?\s*\n?', '', purpose).strip()
    example = re.sub(r'^[Ee]xample\s*[:：]?\s*\n?', '', example).strip()
    pitfall = re.sub(r'^[Pp]itfall\s*[:：]?\s*\n?', '', pitfall).strip()
    pitfall = pitfall.lstrip('⚠').lstrip(': ').strip()

    return {
        'code': code,
        'name': name,
        'formula': formula_text,
        'variables': final_vars,
        'purpose': purpose,
        'example': example,
        'pitfall': pitfall,
        'decision': decision,
        'notes': notes,
    }


def is_formula_table(table):
    if len(table.rows) < 2:
        return False
    first_row_cells = [clean(c.text) for c in table.rows[0].cells]
    if len(first_row_cells) >= 2 and first_row_cells[0] and first_row_cells[0] == first_row_cells[1]:
        if re.match(r'^\d+(\.\d+[a-z]?)?\s+\S', first_row_cells[0]):
            return True
    return False


def is_sustainability_note(table):
    if len(table.rows) == 0:
        return False
    text = clean(table.rows[0].cells[0].text)
    return text.startswith('♻') or 'Sustainability Note' in text


def is_decision_box(table):
    if len(table.rows) == 0:
        return False
    text = clean(table.rows[0].cells[0].text)
    return text.startswith('▲') or text.startswith('Decision') or 'Decision:' in text[:30]


def is_quick_reference(table):
    if len(table.rows) == 0:
        return False
    text = clean(table.rows[0].cells[0].text)
    return text == 'Quick Reference'


def is_summary_table(table):
    """Summary tables have a header row like 'Metric | Formula | Use'."""
    if len(table.rows) == 0:
        return False
    first_row = [clean(c.text) for c in table.rows[0].cells]
    headers = [h.lower() for h in first_row if h]
    if not headers:
        return False
    summary_keywords = ['metric', 'formula', 'index', 'parameter', 'crop', 'species',
                        'bbch', 'month', 'season', 'quantity', 'term', 'variety',
                        'stage', 'soil', 'parameter', 'category', 'soil parameter',
                        'poor', 'good', 'low', 'high']
    return any(any(kw in h for kw in summary_keywords) for h in headers)


def main():
    doc = Document(DOC_PATH)
    body = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables

    items = []
    p_idx = 0
    t_idx = 0
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            items.append(('para', p_idx))
            p_idx += 1
        elif child.tag == qn('w:tbl'):
            items.append(('table', t_idx))
            t_idx += 1

    parts = []
    current_part_obj = None
    current_chapter_obj = None
    current_sub_section_title = None

    for kind, idx in items:
        if kind == 'para':
            para = paragraphs[idx]
            text = clean(para.text)
            if not text:
                continue
            style = para.style.name if para.style else ''

            part_roman = detect_part_number(text)
            if part_roman:
                current_part_obj = {
                    'roman': part_roman,
                    'title': '',
                    'chapters': []
                }
                parts.append(current_part_obj)
                current_chapter_obj = None
                current_sub_section_title = None
                continue

            if current_part_obj and not current_part_obj['title'] and not style.startswith('Heading'):
                current_part_obj['title'] = text
                continue

            if style == 'Heading 1':
                parsed = parse_chapter_title(text)
                if parsed:
                    num, title = parsed
                    current_chapter_obj = {
                        'number': num,
                        'title': title,
                        'intro': '',
                        'sub_sections': [],
                        'formulas': [],
                        'summary_tables': [],
                        'sustainability_notes': [],
                        'decision_boxes': []
                    }
                    if current_part_obj is None:
                        current_part_obj = {'roman': '?', 'title': 'Uncategorized', 'chapters': []}
                        parts.append(current_part_obj)
                    current_part_obj['chapters'].append(current_chapter_obj)
                    current_sub_section_title = None
                continue

            if style == 'Heading 2':
                if current_chapter_obj is not None:
                    current_sub_section_title = text
                    current_chapter_obj['sub_sections'].append({
                        'title': text,
                        'formulas': []
                    })
                continue

            if current_chapter_obj is not None and not current_chapter_obj['formulas'] and not current_chapter_obj['summary_tables']:
                if not style.startswith('Heading'):
                    if current_chapter_obj['intro']:
                        current_chapter_obj['intro'] += '\n' + text
                    else:
                        current_chapter_obj['intro'] = text
        elif kind == 'table':
            table = tables[idx]
            if current_chapter_obj is None:
                continue

            if is_quick_reference(table):
                continue
            if is_sustainability_note(table):
                text = clean(table.rows[0].cells[0].text)
                current_chapter_obj['sustainability_notes'].append(text)
                continue
            if is_decision_box(table):
                text = clean(table.rows[0].cells[0].text)
                current_chapter_obj['decision_boxes'].append(text)
                continue
            if is_formula_table(table):
                formula = parse_formula_table(table)
                if formula:
                    if current_sub_section_title and current_chapter_obj['sub_sections']:
                        current_chapter_obj['sub_sections'][-1]['formulas'].append(formula)
                    else:
                        current_chapter_obj['formulas'].append(formula)
                continue
            # Treat everything else as a summary / reference table
            rows = []
            for r in table.rows:
                cells = [clean(c.text) for c in r.cells]
                uniq = []
                for c in cells:
                    if not uniq or uniq[-1] != c:
                        uniq.append(c)
                rows.append(uniq)
            current_chapter_obj['summary_tables'].append({'title': '', 'rows': rows})

    total_formulas = 0
    for p in parts:
        for c in p['chapters']:
            total_formulas += len(c['formulas'])
            for ss in c['sub_sections']:
                total_formulas += len(ss['formulas'])

    all_formulas = []
    for p in parts:
        for c in p['chapters']:
            for f in c['formulas']:
                all_formulas.append({**f, 'part': p['title'], 'part_roman': p['roman'],
                                     'chapter': c['title'], 'chapter_number': c['number']})
            for ss in c['sub_sections']:
                for f in ss['formulas']:
                    all_formulas.append({**f, 'part': p['title'], 'part_roman': p['roman'],
                                         'chapter': c['title'], 'chapter_number': c['number'],
                                         'sub_section': ss['title']})

    # Build category index for filtering
    categories = {}
    for f in all_formulas:
        cat = f['part']
        if cat not in categories:
            categories[cat] = 0
        categories[cat] += 1

    output = {
        'meta': {
            'title': 'Agri-Formulas & Metrics Handbook',
            'subtitle': 'A Practical Handbook for Crop & Animal Production',
            'version': 'v3 Complete',
            'total_parts': len(parts),
            'total_chapters': sum(len(p['chapters']) for p in parts),
            'total_formulas': total_formulas,
            'categories': categories
        },
        'parts': parts,
        'all_formulas': all_formulas
    }

    with open(OUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"Wrote {OUT_PATH}")
    print(f"Parts: {len(parts)}")
    print(f"Chapters: {sum(len(p['chapters']) for p in parts)}")
    print(f"Formulas: {total_formulas}")
    print(f"All formulas indexed: {len(all_formulas)}")


if __name__ == '__main__':
    main()
